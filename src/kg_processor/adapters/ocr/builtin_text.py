"""Simple text/PDF/Office/HTML parser used for deterministic smoke fixtures."""

from __future__ import annotations

import re
import zipfile
from html.parser import HTMLParser
from pathlib import Path
from xml.etree import ElementTree

from docx import Document as DocxDocument
from pypdf import PdfReader

from kg_processor.domain.documents import InputFile, LayoutBlock, ParsedDocument, ParsedPage
from kg_processor.domain.ids import stable_id
from kg_processor.ports.ocr import OcrOptions


class BuiltinTextOcrProvider:
    """Deterministic local parser for tests and text-native documents.

    This adapter does not attempt image OCR. It is intentionally explicit so
    production jobs can choose MinerU, Tesseract, or Snowflake instead of
    silently falling back to basic text extraction.
    """

    def parse(self, file: InputFile, options: OcrOptions) -> ParsedDocument:
        """Parse text-native formats into deterministic one-block pages."""

        suffix = file.path.suffix.lower()
        if suffix in {".txt", ".md", ".markdown"}:
            pages = [_page_from_text(file, 1, file.path.read_text(encoding="utf-8"))]
        elif suffix == ".pdf":
            pages = self._parse_pdf(file.path, file)
        elif suffix == ".docx":
            pages = self._parse_docx(file.path, file)
        elif suffix in {".html", ".htm"}:
            pages = self._parse_html(file.path, file)
        elif suffix == ".pptx":
            pages = self._parse_pptx(file.path, file)
        elif suffix == ".xlsx":
            pages = self._parse_xlsx(file.path, file)
        else:
            raise ValueError(f"builtin_text OCR does not support {suffix}")
        pages = _select_pages(pages, options.page_range)
        return ParsedDocument(
            file_id=file.id,
            checksum=file.checksum,
            source_uri=file.source_uri,
            mime_type=file.mime_type,
            pages=pages,
            provider_metadata={
                "provider": "builtin_text",
                "language": options.language,
                "page_range": options.page_range,
            },
        )

    def _parse_pdf(self, path: Path, file: InputFile) -> list[ParsedPage]:
        reader = PdfReader(str(path))
        pages: list[ParsedPage] = []
        for index, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            pages.append(_page_from_text(file, index, text))
        return pages

    def _parse_docx(self, path: Path, file: InputFile) -> list[ParsedPage]:
        doc = DocxDocument(str(path))
        text = "\n".join(p.text for p in doc.paragraphs if p.text)
        return [_page_from_text(file, 1, text)]

    def _parse_html(self, path: Path, file: InputFile) -> list[ParsedPage]:
        parser = _VisibleTextHtmlParser()
        parser.feed(path.read_text(encoding="utf-8", errors="replace"))
        return [_page_from_text(file, 1, parser.text())]

    def _parse_pptx(self, path: Path, file: InputFile) -> list[ParsedPage]:
        pages: list[ParsedPage] = []
        with zipfile.ZipFile(path) as archive:
            slide_names = sorted(
                (name for name in archive.namelist() if _is_pptx_slide_xml(name)),
                key=_slide_sort_key,
            )
            for index, name in enumerate(slide_names, start=1):
                text = _pptx_slide_text(archive.read(name))
                pages.append(_page_from_text(file, index, text))
        if not pages:
            raise ValueError(f"PPTX file contains no slides: {path}")
        return pages

    def _parse_xlsx(self, path: Path, file: InputFile) -> list[ParsedPage]:
        pages: list[ParsedPage] = []
        with zipfile.ZipFile(path) as archive:
            shared_strings = _xlsx_shared_strings(archive)
            sheet_names = sorted(
                (name for name in archive.namelist() if _is_xlsx_sheet_xml(name)),
                key=_xlsx_sheet_sort_key,
            )
            for index, name in enumerate(sheet_names, start=1):
                text = _xlsx_sheet_text(archive.read(name), shared_strings)
                pages.append(_page_from_text(file, index, text))
        if not pages:
            raise ValueError(f"XLSX file contains no worksheets: {path}")
        return pages


def _page_from_text(file: InputFile, page_number: int, text: str) -> ParsedPage:
    block = LayoutBlock(
        id=stable_id("block", file.id, page_number, text[:256]),
        page_number=page_number,
        kind="text",
        text=text,
    )
    return ParsedPage(
        page_number=page_number,
        markdown=text,
        raw_text=text,
        blocks=[block],
    )


def _select_pages(pages: list[ParsedPage], page_range: str | None) -> list[ParsedPage]:
    if page_range is None or not page_range.strip():
        return pages
    selected_numbers = _parse_builtin_page_range(page_range)
    selected = [page for page in pages if page.page_number in selected_numbers]
    if not selected:
        raise ValueError(f"builtin_text page_range '{page_range}' selected no pages")
    return selected


def _parse_builtin_page_range(page_range: str) -> set[int]:
    selected: set[int] = set()
    for part in page_range.split(","):
        raw_part = part.strip()
        if not raw_part:
            raise ValueError(f"Invalid builtin_text page_range '{page_range}'")
        if "-" in raw_part:
            start_raw, end_raw = raw_part.split("-", 1)
            start = _parse_builtin_page_number(start_raw, page_range)
            end = _parse_builtin_page_number(end_raw, page_range)
            if end < start:
                raise ValueError(
                    f"Invalid builtin_text page_range '{page_range}': end page is before start page"
                )
            selected.update(range(start, end + 1))
        else:
            selected.add(_parse_builtin_page_number(raw_part, page_range))
    return selected


def _parse_builtin_page_number(value: str, page_range: str) -> int:
    if not value.isdigit() or int(value) < 1:
        raise ValueError(
            f"Invalid builtin_text page number '{value}' in page_range '{page_range}'. "
            "builtin_text page ranges are one-based."
        )
    return int(value)


class _VisibleTextHtmlParser(HTMLParser):
    """Extracts visible text from simple HTML without browser dependencies."""

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self._parts: list[str] = []
        self._ignored_depth = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        """Track ignored blocks and paragraph-like separators."""

        if tag.lower() in {"script", "style", "noscript"}:
            self._ignored_depth += 1
        elif tag.lower() in {"br", "p", "div", "section", "article", "li", "tr", "h1", "h2", "h3"}:
            self._parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        """Close ignored blocks and preserve paragraph-like separators."""

        if tag.lower() in {"script", "style", "noscript"} and self._ignored_depth:
            self._ignored_depth -= 1
        elif tag.lower() in {"p", "div", "section", "article", "li", "tr", "h1", "h2", "h3"}:
            self._parts.append("\n")

    def handle_data(self, data: str) -> None:
        """Capture text nodes that are not inside ignored tags."""

        if not self._ignored_depth:
            self._parts.append(data)

    def text(self) -> str:
        """Return normalized visible text suitable for chunking."""

        # Collapse layout whitespace but preserve paragraph-ish breaks so
        # chunking sees human-readable text rather than a browser DOM dump.
        lines = []
        for line in "".join(self._parts).splitlines():
            normalized = re.sub(r"\s+", " ", line).strip()
            if normalized:
                lines.append(normalized)
        return "\n".join(lines)


def _is_pptx_slide_xml(name: str) -> bool:
    return bool(re.fullmatch(r"ppt/slides/slide\d+\.xml", name))


def _slide_sort_key(name: str) -> int:
    match = re.search(r"slide(\d+)\.xml$", name)
    return int(match.group(1)) if match else 0


def _pptx_slide_text(xml_payload: bytes) -> str:
    root = ElementTree.fromstring(xml_payload)
    values = [
        node.text or ""
        for node in root.iter()
        if node.tag.rsplit("}", 1)[-1] == "t" and node.text
    ]
    return "\n".join(value.strip() for value in values if value.strip())


def _is_xlsx_sheet_xml(name: str) -> bool:
    return bool(re.fullmatch(r"xl/worksheets/sheet\d+\.xml", name))


def _xlsx_sheet_sort_key(name: str) -> int:
    match = re.search(r"sheet(\d+)\.xml$", name)
    return int(match.group(1)) if match else 0


def _xlsx_shared_strings(archive: zipfile.ZipFile) -> list[str]:
    try:
        payload = archive.read("xl/sharedStrings.xml")
    except KeyError:
        return []
    root = ElementTree.fromstring(payload)
    return [
        "".join(_iter_text_nodes(item)).strip()
        for item in root
        if _local_name(item.tag) == "si"
    ]


def _xlsx_sheet_text(xml_payload: bytes, shared_strings: list[str]) -> str:
    root = ElementTree.fromstring(xml_payload)
    rows: list[str] = []
    for row in root.iter():
        if _local_name(row.tag) != "row":
            continue
        values = [
            value
            for cell in row
            if _local_name(cell.tag) == "c"
            for value in [_xlsx_cell_text(cell, shared_strings)]
            if value
        ]
        if values:
            rows.append("\t".join(values))
    return "\n".join(rows)


def _xlsx_cell_text(cell: ElementTree.Element, shared_strings: list[str]) -> str:
    cell_type = cell.attrib.get("t")
    if cell_type == "inlineStr":
        return "".join(_iter_text_nodes(cell)).strip()
    value = next((child.text or "" for child in cell if _local_name(child.tag) == "v"), "")
    if cell_type == "s" and value.isdigit():
        index = int(value)
        if index < len(shared_strings):
            return shared_strings[index]
    return value.strip()


def _iter_text_nodes(node: ElementTree.Element) -> list[str]:
    return [child.text or "" for child in node.iter() if _local_name(child.tag) == "t"]


def _local_name(tag: str) -> str:
    return tag.rsplit("}", 1)[-1]
