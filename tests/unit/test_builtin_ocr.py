from __future__ import annotations

import inspect
from io import BytesIO
from pathlib import Path
from time import perf_counter
from types import SimpleNamespace
from zipfile import ZipFile

import pytest
from defusedxml.common import EntitiesForbidden
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pypdf.errors import LimitReachedError

from kg_processor.adapters.files.local import LocalFileSource
from kg_processor.adapters.ocr import builtin_text
from kg_processor.adapters.ocr.builtin_text import BuiltinTextOcrProvider
from kg_processor.ports.ocr import OcrOptions


def test_builtin_text_ocr_parses_text_fixture() -> None:
    source = LocalFileSource(Path("tests/fixtures/simple.txt"))
    file = source.list_files()[0]

    parsed = BuiltinTextOcrProvider().parse(file, OcrOptions())

    assert parsed.file_id == file.id
    assert parsed.pages[0].page_number == 1
    assert "Alice Smith" in parsed.pages[0].raw_text
    assert parsed.provider_metadata["provider"] == "builtin_text"


def test_builtin_text_ocr_rejects_a_document_with_no_text_layer(tmp_path: Path) -> None:
    """A scanned page has no text layer, and an empty parse is not a success.

    Accepting it records the file as OCR'd while it contributes nothing to the
    graph, which is indistinguishable from a document that genuinely says
    nothing until the graph is inspected.
    """

    text_file = tmp_path / "scanned.txt"
    text_file.write_text("   \n\n\t\n", encoding="utf-8")
    file = LocalFileSource(text_file).list_files()[0]

    with pytest.raises(RuntimeError, match="extracted no text"):
        BuiltinTextOcrProvider().parse(file, OcrOptions())


def test_builtin_text_ocr_decodes_latin1_without_replacement(tmp_path: Path) -> None:
    text_file = tmp_path / "sample.txt"
    text_file.write_bytes("Café résumé naïve".encode("latin-1"))
    file = LocalFileSource(text_file).list_files()[0]

    parsed = BuiltinTextOcrProvider().parse(file, OcrOptions())

    assert parsed.pages[0].raw_text == "Café résumé naïve"
    assert "\ufffd" not in parsed.pages[0].raw_text


def test_builtin_text_ocr_decodes_utf16_bom_without_nul_bytes(tmp_path: Path) -> None:
    text_file = tmp_path / "sample.txt"
    text_file.write_bytes("Hello world café".encode("utf-16"))

    parsed = BuiltinTextOcrProvider().parse(
        LocalFileSource(text_file).list_files()[0], OcrOptions()
    )

    assert parsed.pages[0].raw_text == "Hello world café"
    assert "\x00" not in parsed.pages[0].raw_text


def test_builtin_text_ocr_sniffs_bomless_utf16(tmp_path: Path) -> None:
    text_file = tmp_path / "sample.txt"
    text_file.write_bytes("Hello world café".encode("utf-16-le"))

    parsed = BuiltinTextOcrProvider().parse(
        LocalFileSource(text_file).list_files()[0], OcrOptions()
    )

    assert parsed.pages[0].raw_text == "Hello world café"


def test_builtin_text_ocr_preserves_mixed_utf8_and_cp1252_bytes(tmp_path: Path) -> None:
    text_file = tmp_path / "sample.txt"
    text_file.write_bytes("UTF-8 café — legacy ".encode() + b"\x93quote\x94")

    parsed = BuiltinTextOcrProvider().parse(
        LocalFileSource(text_file).list_files()[0], OcrOptions()
    )

    assert parsed.pages[0].raw_text == "UTF-8 café — legacy “quote”"


def test_mixed_utf8_decoder_scales_linearly_for_legacy_bytes() -> None:
    payload = b"caf\xe9 " * 200_000

    started = perf_counter()
    decoded = builtin_text._decode_mixed_utf8(payload)
    elapsed = perf_counter() - started

    assert decoded.startswith("café café")
    assert len(decoded) == len(payload)
    assert elapsed < 2.0


def test_builtin_text_ocr_strips_utf8_bom_before_heading_classification(tmp_path: Path) -> None:
    markdown = tmp_path / "sample.md"
    markdown.write_bytes("# First heading\nBody".encode("utf-8-sig"))

    parsed = BuiltinTextOcrProvider().parse(LocalFileSource(markdown).list_files()[0], OcrOptions())

    assert parsed.pages[0].raw_text.startswith("# First heading")
    assert parsed.pages[0].blocks[0].kind == "heading"
    assert "\ufeff" not in parsed.pages[0].raw_text


def test_builtin_text_ocr_wraps_corrupt_pdf_errors(tmp_path: Path) -> None:
    pdf = tmp_path / "sample.pdf"
    pdf.write_bytes(b"not a pdf")
    file = LocalFileSource(pdf).list_files()[0]

    with pytest.raises(RuntimeError, match="builtin_text failed to parse PDF"):
        BuiltinTextOcrProvider().parse(file, OcrOptions())


def test_builtin_text_ocr_uses_pdfium_when_primary_pdf_parser_fails(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Preserve physical page text when pypdf rejects unusual font metadata."""

    pdf = tmp_path / "sample.pdf"
    pdf.write_bytes(b"%PDF-1.4\n")
    file = LocalFileSource(pdf).list_files()[0]

    class _FailingReader:
        """Represent a pypdf constructor failure before page extraction."""

        def __init__(self, _path: str) -> None:
            raise RuntimeError("unsupported font descriptor")

    monkeypatch.setattr(builtin_text, "PdfReader", _FailingReader)
    monkeypatch.setattr(
        BuiltinTextOcrProvider,
        "_parse_pdf_with_pdfium",
        lambda _self, _path, _file: [builtin_text._page_from_text(file, 1, "Alice works at Acme.")],
    )

    parsed = BuiltinTextOcrProvider().parse(file, OcrOptions())

    assert [page.page_number for page in parsed.pages] == [1]
    assert parsed.pages[0].raw_text == "Alice works at Acme."


def test_pdfium_fallback_applies_resource_limits_before_opening_pdf() -> None:
    worker_source = inspect.getsource(builtin_text._pdfium_fallback_worker)

    assert worker_source.index("_apply_pdfium_resource_limits()") < worker_source.index(
        "pdfium.PdfDocument(path)"
    )
    assert builtin_text.MAX_PDFIUM_FALLBACK_MEMORY_BYTES > 0
    assert builtin_text.MAX_PDFIUM_FALLBACK_TEXT_BYTES > 0
    assert builtin_text.MAX_PDFIUM_FALLBACK_PAGES > 0


def test_builtin_text_ocr_uses_pdfium_for_primary_value_errors(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    pdf = tmp_path / "sample.pdf"
    pdf.write_bytes(b"%PDF-1.4\n")
    file = LocalFileSource(pdf).list_files()[0]

    class _FailingReader:
        def __init__(self, _path: str) -> None:
            raise ValueError("malformed ASCII85 content")

    monkeypatch.setattr(builtin_text, "PdfReader", _FailingReader)
    monkeypatch.setattr(
        BuiltinTextOcrProvider,
        "_parse_pdf_with_pdfium",
        lambda _self, _path, _file: [builtin_text._page_from_text(file, 1, "recovered")],
    )

    assert BuiltinTextOcrProvider().parse(file, OcrOptions()).pages[0].raw_text == "recovered"


def test_builtin_text_ocr_does_not_fallback_after_pdf_decompression_limit(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    pdf = tmp_path / "bomb.pdf"
    pdf.write_bytes(b"%PDF-1.4\n")
    file = LocalFileSource(pdf).list_files()[0]

    class _LimitedReader:
        def __init__(self, _path: str) -> None:
            raise LimitReachedError("decompression limit")

    monkeypatch.setattr(builtin_text, "PdfReader", _LimitedReader)
    monkeypatch.setattr(
        BuiltinTextOcrProvider,
        "_parse_pdf_with_pdfium",
        lambda *_args: pytest.fail("unsafe PDFium fallback was invoked"),
    )

    with pytest.raises(LimitReachedError):
        BuiltinTextOcrProvider().parse(file, OcrOptions())


def test_office_member_reader_enforces_actual_streamed_output(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    info = SimpleNamespace(file_size=1, filename="word/document.xml")

    class _Archive:
        def getinfo(self, _name: str) -> object:
            return info

        def open(self, _info: object) -> BytesIO:
            return BytesIO(b"oversized")

    monkeypatch.setattr(builtin_text, "MAX_OFFICE_ENTRY_BYTES", 1)

    with pytest.raises(ValueError, match="actual output exceeds"):
        builtin_text._read_zip_member(tmp_path / "bomb.docx", _Archive(), info.filename)  # type: ignore[arg-type]


def test_builtin_text_ocr_parses_docx_body_tables_headers_and_footers(tmp_path: Path) -> None:
    """Protect ordered DOCX extraction across body, tables, headers, and footers.

    Table structure must remain visible as typed rows.
    """

    docx = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("Body paragraph")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Alice Smith"
    table.cell(0, 1).text = "Acme Corp"
    section = document.sections[0]
    section.header.paragraphs[0].text = "Header text"
    section.footer.paragraphs[0].text = "Footer text"
    document.save(str(docx))
    file = LocalFileSource(docx).list_files()[0]

    parsed = BuiltinTextOcrProvider().parse(file, OcrOptions())

    assert "Body paragraph" in parsed.pages[0].raw_text
    assert "Alice Smith | Acme Corp" in parsed.pages[0].raw_text
    assert any(block.kind == "table_row" for block in parsed.pages[0].blocks)
    assert "Header text" in parsed.pages[0].raw_text
    assert "Footer text" in parsed.pages[0].raw_text


def test_builtin_text_ocr_extracts_docx_text_boxes_without_namespace_errors(
    tmp_path: Path,
) -> None:
    docx = tmp_path / "text-box.docx"
    document = Document()
    paragraph = document.add_paragraph("Body text")
    text_box = OxmlElement("w:txbxContent")
    text_box_paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "Text box evidence"
    run.append(text)
    text_box_paragraph.append(run)
    text_box.append(text_box_paragraph)
    paragraph._p.append(text_box)  # noqa: SLF001 - create a real DOCX textbox fixture
    document.save(str(docx))

    parsed = BuiltinTextOcrProvider().parse(LocalFileSource(docx).list_files()[0], OcrOptions())

    assert "Body text" in parsed.pages[0].raw_text
    assert "Text box evidence" in parsed.pages[0].raw_text


def test_builtin_text_ocr_preserves_docx_paragraph_break_and_tab_boundaries(
    tmp_path: Path,
) -> None:
    docx = tmp_path / "boundaries.docx"
    document = Document()
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Revenue"
    table.cell(0, 0).add_paragraph("$1.2M")
    broken = document.add_paragraph()
    broken.add_run("Line one").add_break()
    broken.add_run("Line two")
    tabbed = document.add_paragraph()
    tabbed.add_run("Name").add_tab()
    tabbed.add_run("Value")
    document.save(str(docx))

    parsed = BuiltinTextOcrProvider().parse(LocalFileSource(docx).list_files()[0], OcrOptions())
    text = parsed.pages[0].raw_text

    assert "Revenue\n$1.2M" in text
    assert "Line one\nLine two" in text
    assert "Name | Value" in text
    assert "Revenue$1.2M" not in text
    assert "Line oneLine two" not in text
    assert "NameValue" not in text


def test_builtin_text_ocr_does_not_inherit_heading_style_from_nested_text_box(
    tmp_path: Path,
) -> None:
    docx = tmp_path / "nested-heading.docx"
    document = Document()
    paragraph = document.add_paragraph("Normal body")
    text_box = OxmlElement("w:txbxContent")
    nested_paragraph = OxmlElement("w:p")
    properties = OxmlElement("w:pPr")
    style = OxmlElement("w:pStyle")
    style.set(qn("w:val"), "Heading1")
    properties.append(style)
    nested_paragraph.append(properties)
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "Nested heading"
    run.append(text)
    nested_paragraph.append(run)
    text_box.append(nested_paragraph)
    paragraph._p.append(text_box)  # noqa: SLF001 - create a real DOCX textbox fixture
    document.save(str(docx))

    parsed = BuiltinTextOcrProvider().parse(LocalFileSource(docx).list_files()[0], OcrOptions())

    assert parsed.pages[0].raw_text == "Normal body\nNested heading"
    assert all(block.kind == "paragraph" for block in parsed.pages[0].blocks)


def test_builtin_text_ocr_extracts_docx_content_controls(tmp_path: Path) -> None:
    docx = tmp_path / "content-control.docx"
    document = Document()
    document.add_paragraph("Body")
    content_control = OxmlElement("w:sdt")
    content = OxmlElement("w:sdtContent")
    paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "Controlled title"
    run.append(text)
    paragraph.append(run)
    content.append(paragraph)
    content_control.append(content)
    document._element.body.insert(0, content_control)  # noqa: SLF001 - real SDT fixture
    document.save(str(docx))

    parsed = BuiltinTextOcrProvider().parse(LocalFileSource(docx).list_files()[0], OcrOptions())

    assert parsed.pages[0].raw_text.startswith("Controlled title\nBody")


def test_builtin_text_ocr_bounds_docx_xml_before_python_docx_loads_it(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    docx = tmp_path / "oversized.docx"
    document = Document()
    document.add_paragraph("Small source")
    document.save(str(docx))
    monkeypatch.setattr(builtin_text, "MAX_OFFICE_ENTRY_BYTES", 1)

    with pytest.raises(ValueError, match="above the 1 byte limit"):
        BuiltinTextOcrProvider().parse(LocalFileSource(docx).list_files()[0], OcrOptions())


def test_builtin_text_ocr_ignores_oversized_unread_docx_media(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    docx = tmp_path / "media.docx"
    document = Document()
    document.add_paragraph("Readable body")
    document.save(str(docx))
    with ZipFile(docx) as archive:
        document_size = archive.getinfo("word/document.xml").file_size
    entry_limit = document_size + 1
    with ZipFile(docx, "a") as archive:
        archive.writestr("word/media/large-video.bin", b"x" * (entry_limit + 1))
    monkeypatch.setattr(builtin_text, "MAX_OFFICE_ENTRY_BYTES", entry_limit)

    parsed = BuiltinTextOcrProvider().parse(LocalFileSource(docx).list_files()[0], OcrOptions())

    assert "Readable body" in parsed.pages[0].raw_text


def test_builtin_text_ocr_bounds_aggregate_docx_xml(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    docx = tmp_path / "aggregate.docx"
    document = Document()
    document.add_paragraph("Body")
    document.save(str(docx))
    header_xml = (
        '<w:hdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        "<w:p><w:r><w:t>Header content</w:t></w:r></w:p></w:hdr>"
    )
    with ZipFile(docx, "a") as archive:
        archive.writestr("word/header1.xml", header_xml)
        archive.writestr("word/header2.xml", header_xml)
    with ZipFile(docx) as archive:
        readable_sizes = [
            archive.getinfo(name).file_size
            for name in ("word/document.xml", "word/header1.xml", "word/header2.xml")
        ]
    monkeypatch.setattr(builtin_text, "MAX_OFFICE_ENTRY_BYTES", max(readable_sizes) + 1)
    monkeypatch.setattr(builtin_text, "MAX_OFFICE_TOTAL_BYTES", sum(readable_sizes) - 1)

    with pytest.raises(ValueError, match="aggregate limit"):
        BuiltinTextOcrProvider().parse(LocalFileSource(docx).list_files()[0], OcrOptions())


def test_builtin_text_ocr_parses_html(tmp_path: Path) -> None:
    """Verify visible HTML structure survives while script and style content is excluded."""

    html = tmp_path / "sample.html"
    html.write_text(
        """
<html>
  <head><style>.hidden { color: red }</style><script>ignored()</script></head>
  <body>
    <h1>Alice Smith</h1><p>Alice works at Acme Corp.</p>
    <table><tr><th>System</th><th>Place</th></tr><tr><td>Judo</td><td>Tokyo</td></tr></table>
  </body>
</html>
""",
        encoding="utf-8",
    )
    file = LocalFileSource(html).list_files()[0]

    parsed = BuiltinTextOcrProvider().parse(file, OcrOptions())

    assert "Alice Smith" in parsed.pages[0].raw_text
    assert "Alice works at Acme Corp." in parsed.pages[0].raw_text
    assert "System | Place" in parsed.pages[0].raw_text
    assert "Judo | Tokyo" in parsed.pages[0].raw_text
    assert [block.kind for block in parsed.pages[0].blocks].count("table_row") == 2
    assert parsed.pages[0].blocks[0].kind == "heading"
    assert "ignored" not in parsed.pages[0].raw_text


def test_builtin_text_ocr_uses_html_meta_charset(tmp_path: Path) -> None:
    html = tmp_path / "latin1.html"
    html.write_bytes(
        '<html><head><meta charset="iso-8859-1"></head><body>Café résumé</body></html>'.encode(
            "latin-1"
        )
    )

    parsed = BuiltinTextOcrProvider().parse(LocalFileSource(html).list_files()[0], OcrOptions())

    assert parsed.pages[0].raw_text == "Café résumé"
    assert "\ufffd" not in parsed.pages[0].raw_text


def test_builtin_text_ocr_sniffs_only_real_meta_charset_tags(tmp_path: Path) -> None:
    html = tmp_path / "script-meta.html"
    html.write_bytes(
        b'<script>const example = "<meta charset=unicode>";</script>'
        b'<meta data-origin="export" charset="utf-8"><p>Caf\xc3\xa9</p>'
    )

    parsed = BuiltinTextOcrProvider().parse(LocalFileSource(html).list_files()[0], OcrOptions())

    assert parsed.pages[0].raw_text == "Café"


def test_builtin_text_ocr_parses_pptx_slides(tmp_path: Path) -> None:
    pptx = tmp_path / "sample.pptx"
    _write_minimal_pptx(pptx)
    file = LocalFileSource(pptx).list_files()[0]

    parsed = BuiltinTextOcrProvider().parse(file, OcrOptions())

    assert [page.page_number for page in parsed.pages] == [1, 2]
    assert "Alice Smith" in parsed.pages[0].raw_text
    assert "Acme Corp" in parsed.pages[1].raw_text


def test_builtin_text_ocr_uses_presentation_relationship_order_and_part_names(
    tmp_path: Path,
) -> None:
    pptx = tmp_path / "relationships.pptx"
    slide_xml = """
<p:sld xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"
       xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
  <p:cSld><p:spTree><p:sp><p:txBody><a:p><a:r><a:t>{text}</a:t></a:r></a:p>
  </p:txBody></p:sp></p:spTree></p:cSld>
</p:sld>
"""
    with ZipFile(pptx, "w") as archive:
        archive.writestr(
            "ppt/presentation.xml",
            """
<p:presentation xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"
 xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
 <p:sldIdLst><p:sldId id="256" r:id="rSecond"/><p:sldId id="257" r:id="rFirst"/></p:sldIdLst>
</p:presentation>
""",
        )
        archive.writestr(
            "ppt/_rels/presentation.xml.rels",
            """
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
 <Relationship Id="rFirst"
  Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slide"
  Target="custom/alpha.xml"/>
 <Relationship Id="rSecond"
  Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slide"
  Target="custom/zeta.xml"/>
</Relationships>
""",
        )
        archive.writestr("ppt/custom/alpha.xml", slide_xml.format(text="Alpha"))
        archive.writestr("ppt/custom/zeta.xml", slide_xml.format(text="Zeta"))

    parsed = BuiltinTextOcrProvider().parse(LocalFileSource(pptx).list_files()[0], OcrOptions())

    assert [page.raw_text for page in parsed.pages] == ["Zeta", "Alpha"]


def test_builtin_text_ocr_rejects_large_office_zip_entries(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    pptx = tmp_path / "sample.pptx"
    _write_minimal_pptx(pptx)
    file = LocalFileSource(pptx).list_files()[0]
    monkeypatch.setattr(builtin_text, "MAX_OFFICE_ENTRY_BYTES", 1)

    with pytest.raises(ValueError, match="above the 1 byte limit"):
        BuiltinTextOcrProvider().parse(file, OcrOptions())


@pytest.mark.parametrize("suffix", [".pptx", ".xlsx"])
def test_builtin_text_ocr_bounds_aggregate_pptx_and_xlsx_xml(
    suffix: str,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    office_file = tmp_path / f"aggregate{suffix}"
    if suffix == ".pptx":
        _write_minimal_pptx(office_file)
    else:
        _write_minimal_xlsx(office_file)
    with ZipFile(office_file) as archive:
        readable_sizes = [
            info.file_size
            for info in archive.infolist()
            if (
                info.filename == "xl/sharedStrings.xml"
                or info.filename.startswith("xl/worksheets/sheet")
                or info.filename.startswith("ppt/slides/slide")
            )
        ]
    monkeypatch.setattr(builtin_text, "MAX_OFFICE_ENTRY_BYTES", max(readable_sizes) + 1)
    monkeypatch.setattr(builtin_text, "MAX_OFFICE_TOTAL_BYTES", sum(readable_sizes) - 1)

    with pytest.raises(ValueError, match="aggregate limit"):
        BuiltinTextOcrProvider().parse(LocalFileSource(office_file).list_files()[0], OcrOptions())


def test_builtin_text_ocr_uses_hardened_xml_parser_for_pptx(tmp_path: Path) -> None:
    pptx = tmp_path / "entity.pptx"
    with ZipFile(pptx, "w") as archive:
        archive.writestr(
            "ppt/slides/slide1.xml",
            """
<!DOCTYPE x [<!ENTITY boom "expanded">]>
<p:sld xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"
       xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
  <p:cSld>
    <p:spTree>
      <p:sp>
        <p:txBody><a:p><a:r><a:t>&boom;</a:t></a:r></a:p></p:txBody>
      </p:sp>
    </p:spTree>
  </p:cSld>
</p:sld>
""",
        )
    file = LocalFileSource(pptx).list_files()[0]

    with pytest.raises(EntitiesForbidden):
        BuiltinTextOcrProvider().parse(file, OcrOptions())


def test_builtin_text_ocr_filters_pages_with_one_based_page_range(tmp_path: Path) -> None:
    pptx = tmp_path / "sample.pptx"
    _write_minimal_pptx(pptx)
    file = LocalFileSource(pptx).list_files()[0]

    parsed = BuiltinTextOcrProvider().parse(file, OcrOptions(page_range="2"))

    assert [page.page_number for page in parsed.pages] == [2]
    assert "Acme Corp" in parsed.pages[0].raw_text
    assert parsed.provider_metadata["page_range"] == "2"


def test_builtin_text_ocr_rejects_invalid_page_range(tmp_path: Path) -> None:
    pptx = tmp_path / "sample.pptx"
    _write_minimal_pptx(pptx)
    file = LocalFileSource(pptx).list_files()[0]

    with pytest.raises(ValueError, match="one-based"):
        BuiltinTextOcrProvider().parse(file, OcrOptions(page_range="0"))


def test_builtin_text_ocr_parses_xlsx_worksheets(tmp_path: Path) -> None:
    """Verify XLSX worksheets become ordered pages with typed table-row metadata.

    Shared strings and numeric cells must resolve correctly.
    """

    workbook = tmp_path / "sample.xlsx"
    _write_minimal_xlsx(workbook)
    file = LocalFileSource(workbook).list_files()[0]

    parsed = BuiltinTextOcrProvider().parse(file, OcrOptions())

    assert [page.page_number for page in parsed.pages] == [1, 2]
    assert "Alice Smith | Acme Corp" in parsed.pages[0].raw_text
    assert parsed.pages[0].blocks[0].metadata["cells"] == ["Alice Smith", "Acme Corp"]
    assert "Copenhagen | 42" in parsed.pages[1].raw_text
    assert "99" not in parsed.pages[0].raw_text


def test_builtin_text_ocr_uses_workbook_relationship_order_and_part_names(
    tmp_path: Path,
) -> None:
    workbook = tmp_path / "relationships.xlsx"
    worksheet_xml = """
<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
 <sheetData><row r="1"><c r="A1" t="s"><v>{index}</v></c></row></sheetData>
</worksheet>
"""
    with ZipFile(workbook, "w") as archive:
        archive.writestr(
            "xl/workbook.xml",
            """
<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main"
 xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
 <sheets>
  <sheet name="Second" sheetId="2" r:id="rSecond"/>
  <sheet name="First" sheetId="1" r:id="rFirst"/>
 </sheets>
</workbook>
""",
        )
        archive.writestr(
            "xl/_rels/workbook.xml.rels",
            """
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
 <Relationship Id="rFirst"
  Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet"
  Target="custom/alpha.xml"/>
 <Relationship Id="rSecond"
  Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet"
  Target="custom/zeta.xml"/>
 <Relationship Id="rStrings"
  Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/sharedStrings"
  Target="strings/custom.xml"/>
</Relationships>
""",
        )
        archive.writestr(
            "xl/strings/custom.xml",
            """<sst xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
<si><t>Alpha</t></si><si><t>Zeta</t></si></sst>""",
        )
        archive.writestr("xl/custom/alpha.xml", worksheet_xml.format(index=0))
        archive.writestr("xl/custom/zeta.xml", worksheet_xml.format(index=1))

    parsed = BuiltinTextOcrProvider().parse(LocalFileSource(workbook).list_files()[0], OcrOptions())

    assert [page.raw_text for page in parsed.pages] == ["Zeta", "Alpha"]


def test_builtin_text_ocr_excludes_xlsx_phonetic_runs(tmp_path: Path) -> None:
    workbook = tmp_path / "phonetic.xlsx"
    with ZipFile(workbook, "w") as archive:
        archive.writestr(
            "xl/sharedStrings.xml",
            """
<sst xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
  <si><t>東京</t><rPh sb="0" eb="2"><t>とうきょう</t></rPh></si>
</sst>
""",
        )
        archive.writestr(
            "xl/worksheets/sheet1.xml",
            """
<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
  <sheetData><row r="1"><c r="A1" t="s"><v>0</v></c></row></sheetData>
</worksheet>
""",
        )

    parsed = BuiltinTextOcrProvider().parse(LocalFileSource(workbook).list_files()[0], OcrOptions())

    assert parsed.pages[0].raw_text == "東京"
    assert "とうきょう" not in parsed.pages[0].raw_text


def _write_minimal_pptx(path: Path) -> None:
    with ZipFile(path, "w") as archive:
        archive.writestr(
            "ppt/slides/slide2.xml",
            """
<p:sld xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"
       xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
  <p:cSld>
    <p:spTree>
      <p:sp>
        <p:txBody><a:p><a:r><a:t>Acme Corp</a:t></a:r></a:p></p:txBody>
      </p:sp>
    </p:spTree>
  </p:cSld>
</p:sld>
""",
        )
        archive.writestr(
            "ppt/slides/slide1.xml",
            """
<p:sld xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"
       xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
  <p:cSld>
    <p:spTree>
      <p:sp>
        <p:txBody><a:p><a:r><a:t>Alice Smith</a:t></a:r></a:p></p:txBody>
      </p:sp>
    </p:spTree>
  </p:cSld>
</p:sld>
""",
        )


def _write_minimal_xlsx(path: Path) -> None:
    with ZipFile(path, "w") as archive:
        archive.writestr(
            "xl/sharedStrings.xml",
            """
<sst xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
  <si><t>Alice Smith</t></si>
  <si><t>Acme Corp</t></si>
  <si><t>Copenhagen</t></si>
</sst>
""",
        )
        archive.writestr(
            "xl/worksheets/sheet2.xml",
            """
<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
  <sheetData>
    <row r="1"><c r="A1" t="s"><v>2</v></c><c r="B1"><v>42</v></c></row>
  </sheetData>
</worksheet>
""",
        )
        archive.writestr(
            "xl/worksheets/sheet1.xml",
            """
<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
  <sheetData>
    <row r="1">
      <c r="A1" t="s"><v>0</v></c>
      <c r="B1" t="s"><v>1</v></c>
      <c r="C1" t="s"><v>99</v></c>
    </row>
  </sheetData>
</worksheet>
""",
        )
